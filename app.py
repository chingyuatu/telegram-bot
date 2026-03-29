import os
import logging
import requests
import io
from flask import Flask, request
from groq import Groq
from docx import Document
from datetime import datetime

# 1. Logging and App Setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# 2. Environment Variables (Set these in Render Dashboard)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
TELEGRAM_API_URL = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/"

# Initialize Groq Client
client = Groq(api_key=GROQ_API_KEY)

# 3. Core Function: Academic Research Synthesis
def generate_research_report():
    """Uses Groq Llama 3.3 to synthesize latest research insights."""
    research_prompt = (
        "Act as a professional medical researcher. Search and summarize the 3 most significant "
        "academic findings regarding 'Postbiotics and Infant Health' published in 2025-2026. "
        "For each, provide: 1. Study Title, 2. Key Mechanism/Finding, 3. Practical Clinical Implication."
    )
    
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": research_prompt}],
            max_tokens=2000
        )
        return completion.choices[0].message.content
    except Exception as e:
        logger.error(f"Groq API Error: {e}")
        return "Error: Unable to synthesize research data at this time."

# 4. Core Function: Generate Word Document
def create_word_report(content):
    """Generates a .docx file in memory from the research content."""
    doc = Document()
    doc.add_heading('Daily Postbiotics Research Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_heading('Key Research Summaries', level=1)
    doc.add_paragraph(content)
    
    # Save to memory buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 5. Telegram Messaging Helpers
def send_telegram_text(chat_id, text):
    url = f"{TELEGRAM_API_URL}sendMessage"
    payload = {"chat_id": chat_id, "text": text}
    requests.post(url, json=payload, timeout=10)

def send_telegram_document(chat_id, file_stream, caption):
    url = f"{TELEGRAM_API_URL}sendDocument"
    files = {
        'document': ('Postbiotics_Research_Report.docx', file_stream, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    }
    data = {'chat_id': chat_id, 'caption': caption}
    requests.post(url, files=files, data=data, timeout=20)

# 6. Webhook Endpoint
@app.route("/webhook", methods=["POST"])
def webhook():
    data = request.json
    if not data:
        return "No Data", 400

    if "message" in data and "text" in data["message"]:
        chat_id = data["message"]["chat"]["id"]
        incoming
